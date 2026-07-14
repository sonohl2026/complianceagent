import io

import docx

from app.services.parsing.base import Block, ParsedDocument, ParsingError


def parse_docx(content: bytes) -> ParsedDocument:
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise ParsingError(f"Failed to open DOCX: {exc}") from exc

    blocks: list[Block] = []
    title = None

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name if para.style else "") or ""
        heading_level = 0
        if style_name.lower().startswith("title") and title is None:
            title = text
            heading_level = 1
        elif style_name.lower().startswith("heading"):
            digits = "".join(ch for ch in style_name if ch.isdigit())
            heading_level = int(digits) if digits else 1
            heading_level = min(heading_level, 6)
        blocks.append(Block(text=text, page_number=None, heading_level=heading_level))

    for table in document.tables:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        if rows_text:
            blocks.append(Block(text="\n".join(rows_text), page_number=None, heading_level=0))

    if not blocks:
        raise ParsingError("DOCX contains no extractable text")

    return ParsedDocument(title=title, blocks=blocks)
